from pydoc import doc
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import pandas as pd
from app.main import process_results
from docx import Document  # type: ignore[import]
from docx.oxml import OxmlElement  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]
import json
from datetime import datetime
from werkzeug.utils import secure_filename
from app.extractors.pdf_extractor import extract_pdf_data


from app.services.performance_utils import (
    measure_performance,
    ResultCache,
    get_file_stats,
    get_file_hash,
    logger,
)
from app.services.worker_pool import run_in_background

app = Flask(__name__)

# --- IN-MEMORY DUPLICATE TRACKER ---
PROCESSED_HASHES = set()

# --- FEATURE FLAGS ---
USE_ASYNC = os.environ.get("USE_ASYNC", "False").lower() == "true"
ENABLE_CACHE = os.environ.get("ENABLE_CACHE", "True").lower() == "true"

ALLOWED_ORIGINS = [
    "http://localhost:8080",
    "http://127.0.0.1:8080",
    "https://sims-result.netlify.app/",
    "http://localhost:5173",
]

# Allow overriding origins via environment variable
extra_origins = os.environ.get("ORIGINS")
if extra_origins:
    ALLOWED_ORIGINS.extend(extra_origins.split(","))

CORS(
    app,
    resources={
        r"/*": {
            "origins": ALLOWED_ORIGINS,
            "methods": ["GET", "POST", "OPTIONS", "DELETE"],
            "allow_headers": ["Content-Type", "Authorization"],
        }
    },
)


UPLOAD_FOLDER = "uploads"
EXPORT_FOLDER = "outputs"
TEMPLATE_FOLDER = os.path.join("app", "templates")


os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(EXPORT_FOLDER, exist_ok=True)


@app.route("/")
@measure_performance
def home():
    return jsonify(
        {
            "message": "Result Engine API Running",
            "status": "Healthy",
            "cache_enabled": ENABLE_CACHE,
            "async_enabled": USE_ASYNC,
        }
    )


# --- CASTE-BASED FILTERING SYSTEM ---
ALLOWED_CASTES = ["General", "OBC", "SC", "ST"]


@app.route("/upload-caste-filter", methods=["POST"])
def upload_caste_filter():
    file = request.files.get("file")
    caste = request.form.get("caste")

    # Validation
    if not file or not file.filename.lower().endswith(".pdf"):
        return jsonify({"success": False, "message": "Only PDF files are allowed"}), 400

    if not caste or caste not in ALLOWED_CASTES:
        return jsonify(
            {
                "success": False,
                "message": f"Valid Caste category is required {ALLOWED_CASTES}",
            }
        ), 400

    # Secure Storage with Timestamp
    timestamp = int(time.time())
    file_id = secure_filename(file.filename)
    filename = f"result_{timestamp}_{file_id}"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    try:
        # Process and filter the results
        filtered_data = process_caste_result(filepath, caste)

        return jsonify(
            {
                "success": True,
                "message": "File uploaded and processed successfully",
                "data": filtered_data,
                "fileName": filename,
                "caste": caste,
            }
        )
    except Exception as e:
        logger.error(f"Processing error: {str(e)}")
        return jsonify(
            {"success": False, "message": f"Processing error: {str(e)}"}
        ), 500


def process_caste_result(file_path, target_caste):
    """
    Extracts data from the PDF and applies caste-based filtering.
    """
    # Use existing extraction logic
    students = extract_pdf_data(file_path)

    processed_results = []

    for student in students:
        # student objects from extract_pdf_data already have USN, Name and Marks.
        # We need to map them to a caste.
        # In this standalone feature, we will use a deterministic mock based on the USN suffix for demo purposes,
        # or you can integrate your DB lookup here.

        # Simple deterministic mapping for demo/integration
        usn_id = int("".join(filter(str.isdigit, student.usn)) or "0")
        student_caste = ALLOWED_CASTES[usn_id % len(ALLOWED_CASTES)]

        if student_caste.upper() == target_caste.upper():
            processed_results.append(
                {
                    "name": student.name,
                    "usn": student.usn,
                    "overall_total": student.overall_total,
                    "overall_max": student.overall_max,
                    "percentage": round(student.percentage, 2),
                    "result": student.result,
                }
            )

    return {"total_filtered": len(processed_results), "students": processed_results}


@app.route("/files", methods=["GET"])
def list_files():
    files = os.listdir(UPLOAD_FOLDER)
    pdf_files = [f for f in files if f.endswith(".pdf")]
    return jsonify({"success": True, "files": pdf_files})


@app.route("/files/<filename>", methods=["DELETE"])
def delete_file(filename):
    filepath = os.path.join(UPLOAD_FOLDER, secure_filename(filename))
    if os.path.exists(filepath):
        os.remove(filepath)
        return jsonify({"success": True, "message": f"Deleted {filename}"})
    return jsonify({"success": False, "message": "File not found"}), 404


# ===================== UPLOAD =====================
@app.route("/upload", methods=["POST"])
@measure_performance
def upload():
    marks_file = request.files.get("marks")
    caste_file = request.files.get("caste")
    ui_meta = request.form.get("ui_meta")

    if not marks_file or not caste_file:
        return jsonify({"error": "Both files required"}), 400

    ui_meta = json.loads(ui_meta) if ui_meta else None

    # Performance: Only read bytes once
    marks_bytes = marks_file.read()
    caste_bytes = caste_file.read()

    # Combined hash for cache lookup (Instant retrieval for same files)
    current_hash = get_file_hash(marks_bytes + caste_bytes)

    # Save to disk only if necessary (or keep internal)
    marks_filename = f"marks_{current_hash[:8]}.pdf"
    caste_filename = f"caste_{current_hash[:8]}.xlsx"
    marks_path = os.path.join(UPLOAD_FOLDER, marks_filename)
    caste_path = os.path.join(UPLOAD_FOLDER, caste_filename)

    if not os.path.exists(marks_path):
        with open(marks_path, "wb") as f:
            f.write(marks_bytes)
    if not os.path.exists(caste_path):
        with open(caste_path, "wb") as f:
            f.write(caste_bytes)

    app.config["MARKS_PATH"] = marks_path
    app.config["CASTE_PATH"] = caste_path
    app.config["UI_META"] = ui_meta

    # CORE DATA EXTRACTION (Cached internally within process_results)
    results = process_results(marks_path, caste_path, ui_meta)

    return jsonify(results)


# ##======= GENERATE DOCX =====================
@app.route("/generate-report", methods=["GET", "POST"])
@app.route("/download", methods=["GET", "POST"])
def generate_doc_report():

    if "MARKS_PATH" not in app.config:
        return jsonify({"error": "No processed data available"}), 400

    # Backward compatible: when missing, behave exactly like existing implementation.
    # Supported: ?format=internal (default), ?format=public
    download_format = (request.args.get("format") or "internal").strip().lower()
    if download_format not in {"internal", "public"}:
        download_format = "internal"

    # Prefer UI meta sent with this request (POST), fall back to stored one from upload
    ui_meta = app.config.get("UI_META")

    if request.method == "POST":
        payload = request.get_json(silent=True) or {}
        incoming_meta = payload.get("ui_meta")
        print("Incoming ui_meta from UI:", incoming_meta)
        if incoming_meta:
            ui_meta = incoming_meta
            app.config["UI_META"] = ui_meta

    print("Using ui_meta in generate-report:", ui_meta)

    data = process_results(app.config["MARKS_PATH"], app.config["CASTE_PATH"], ui_meta)

    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(
        BASE_DIR, "templates", "Approved Result Analysis Tabulation 2025-26.docx"
    )
    output_path = os.path.join(EXPORT_FOLDER, "Approved_Result_Analysis.docx")

    print("Template path:", template_path)
    print("Exists:", os.path.exists(template_path))
    doc = Document(template_path)

    def _enforce_read_only(doc_obj):
        # "Read-only" without password (removable but meets the "if possible" requirement)
        try:
            settings_el = doc_obj.settings.element
            existing = settings_el.xpath("./w:documentProtection")
            dp = existing[0] if existing else OxmlElement("w:documentProtection")
            dp.set(qn("w:edit"), "readOnly")
            dp.set(qn("w:enforcement"), "1")
            if not existing:
                settings_el.append(dp)
        except Exception:
            pass

    _enforce_read_only(doc)

    metadata = data["metadata"]
    summary = data["summary"]
    rankers = data["rankers"]
    subjects = data["subjects"]
    demographics = data["demographics"]
    centum = data["centum"]

    # ============ HEADER FILL ============

    def replace_header_in_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                text = run.text
                if "Academic Year" in text:
                    run.text = f"Academic Year: {metadata.get('academic_year', '')}"
                if "Name of the Program" in text:
                    run.text = f"Name of the Program: {metadata.get('department', '')}"
                if "BU Examination month" in text:
                    run.text = f"BU Examination month & year: {metadata.get('exam_session', '')}"
                if "Semester" in text:
                    run.text = f"Semester: {metadata.get('semester', '')}"
                if "Date of Declaration" in text:
                    run.text = f"Date of Declaration of Result: {metadata.get('result_date', '')}"

    # Replace in document body

    raw_date = metadata.get("result_date", "")
    formatted_date = raw_date
    if raw_date and "-" in raw_date:
        try:
            dt = datetime.strptime(raw_date, "%Y-%m-%d")
            formatted_date = dt.strftime("%d/%m/%Y")
        except:
            pass

    # Exact labels from USER requirements
    header_mapping = {
        "Academic Year": metadata.get("academic_year", ""),
        "Name of the Program": metadata.get("department", ""),
        "BU Examination month": metadata.get("exam_session", ""),
        "Semester": metadata.get("semester", ""),
        "Date of Declaration": formatted_date,
    }

    def replace_header_in_paragraphs(paragraphs):
        for para in paragraphs:
            para_text = para.text
            for label, value in header_mapping.items():
                if label.lower() in para_text.lower():
                    # Found a paragraph with the target label.
                    label_seen = False
                    for run in para.runs:
                        # Case A: Label and colon are in this run
                        low_text = run.text.lower()
                        if label.lower() in low_text:
                            label_seen = True
                            if ":" in run.text:
                                # Replace everything after the FIRST colon in this run
                                parts = run.text.split(":", 1)
                                run.text = f"{parts[0]}: {value}"
                                # We've filled it, but we might need to clear subsequent underscores
                                continue

                        # Case B: This run contains placeholders (underscores) following a label
                        if label_seen and ("___" in run.text or run.text.strip() == ""):
                            # If we haven't put the value in yet, or this is just extra padding
                            if (
                                value and label.lower() in para_text.lower()
                            ):  # double check
                                # If the previous run didn't already contain the value
                                if value not in para.text:
                                    run.text = value
                                else:
                                    run.text = ""
                            else:
                                run.text = ""
                        elif label_seen and run.text.strip().startswith(":"):
                            # Handle colon in its own run
                            run.text = f": {value}"
                            label_seen = True  # keep looking to clear extra underscores

    # Replace in document body, headers, and all tables

    replace_header_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_header_in_paragraphs(cell.paragraphs)

    # Also replace in page headers (often where this academic info is placed)
    for section in doc.sections:
        replace_header_in_paragraphs(section.header.paragraphs)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_header_in_paragraphs(cell.paragraphs)

    # ============ SUMMARY TABLE ============
    table = doc.tables[0]
    boys = summary["Boys"]
    girls = summary["Girls"]
    total = summary["Total"]

    def fill_summary_row(row_index, block):
        table.rows[row_index].cells[1].text = str(block["appeared"])
        table.rows[row_index].cells[2].text = str(block["distinction"])
        table.rows[row_index].cells[3].text = str(block["first"])
        table.rows[row_index].cells[4].text = str(block["second"])
        table.rows[row_index].cells[5].text = str(block["pass_class"])
        table.rows[row_index].cells[6].text = str(block["passed"])
        table.rows[row_index].cells[7].text = str(block["failed"])
        table.rows[row_index].cells[8].text = f"{round(block['pass_percentage'], 2)}%"

    fill_summary_row(1, boys)
    fill_summary_row(2, girls)
    fill_summary_row(3, total)

    # ============ TOP RANKERS ============
    rank_table = doc.tables[1]
    for i in range(min(3, len(rankers))):
        rank_table.rows[i + 1].cells[1].text = rankers[i]["name"]
        rank_table.rows[i + 1].cells[2].text = rankers[i]["registrationNo"]
        rank_table.rows[i + 1].cells[3].text = str(rankers[i]["marksObtained"])
        rank_table.rows[i + 1].cells[4].text = f"{rankers[i]['percentage']}%"

    # ============ SUBJECT TABLE ============
    subject_table = doc.tables[2]
    for i, sub in enumerate(subjects):
        if i + 1 >= len(subject_table.rows):
            break
        row = subject_table.rows[i + 1].cells
        row[0].text = str(sub["slNo"])
        row[1].text = sub["subject"]
        row[2].text = sub["section"]
        row[3].text = sub["faculty"]
        row[4].text = str(sub["passed"])
        row[5].text = str(sub["failed"])
        row[6].text = str(sub["absent"])
        row[7].text = str(sub["centum"])
        row[8].text = f"{sub['passPercent']}%"
        row[9].text = f"{sub['topper']}%"

    # ============ DEMOGRAPHICS TABLE ============
    demo_table = doc.tables[3]

    appeared = demographics["appeared"]
    passed = demographics["passed"]
    passed_60 = demographics["passed_60"]

    def fill_section(section_title, block_data):
        section_row_index = None
        for i, row in enumerate(demo_table.rows):
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            if section_title in row_text:
                section_row_index = i
                break
        if section_row_index is None:
            return
        for i in range(section_row_index + 1, len(demo_table.rows)):
            if demo_table.rows[i].cells[2].text.strip() == "Total":
                row = demo_table.rows[i]
                g_m = block_data.get("GENERAL", {}).get("MALE", 0)
                g_f = block_data.get("GENERAL", {}).get("FEMALE", 0)

                sc_m = block_data.get("SC", {}).get("MALE", 0)
                sc_f = block_data.get("SC", {}).get("FEMALE", 0)

                st_m = block_data.get("ST", {}).get("MALE", 0)
                st_f = block_data.get("ST", {}).get("FEMALE", 0)

                obc_m = block_data.get("OBC", {}).get("MALE", 0)
                obc_f = block_data.get("OBC", {}).get("FEMALE", 0)

                # ---- Fill GENERAL ----
                row.cells[3].text = str(g_m)
                row.cells[4].text = str(g_f)
                row.cells[5].text = "0"

                # ---- Skip EWS (leave blank) ----

                # ---- Fill SC ----
                row.cells[9].text = str(sc_m)
                row.cells[10].text = str(sc_f)
                row.cells[11].text = "0"

                # ---- Fill ST ----
                row.cells[12].text = str(st_m)
                row.cells[13].text = str(st_f)
                row.cells[14].text = "0"

                # ---- Fill OBC ----
                row.cells[15].text = str(obc_m)
                row.cells[16].text = str(obc_f)
                row.cells[17].text = "0"

                # ---- Fill TOTAL ----
                total_m = g_m + sc_m + st_m + obc_m
                total_f = g_f + sc_f + st_f + obc_f

                row.cells[18].text = str(total_m)
                row.cells[19].text = str(total_f)
                row.cells[20].text = "0"

                break

    # Apply to all 3 blocks (internal only)
    if download_format != "public":
        fill_section("Total Number of Students Appeared", appeared)
        fill_section("Total Number of Students Passed", passed)
        fill_section("Out of Total", passed_60)
    else:
        # PUBLIC: keep the section + table structure, but render the entire
        # demographics table with empty numeric cells (no 0s, no totals).
        try:
            for row in demo_table.rows:
                for cell in row.cells:
                    if any(ch.isdigit() for ch in cell.text):
                        cell.text = ""
        except Exception:
            pass

        # If the template contains any explicit caste labels in this section,
        # remove them as well (avoid "Caste:" with blank values).
        def scrub_caste_labels(paragraphs):
            for para in paragraphs:
                for run in para.runs:
                    if "caste" in run.text.lower():
                        run.text = ""

        scrub_caste_labels(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    scrub_caste_labels(cell.paragraphs)

    # ============ CENTUM ============
    centum_table = doc.tables[-1]

    for i, c in enumerate(centum):
        if i + 1 >= len(centum_table.rows):
            break

        row = centum_table.rows[i + 1].cells
        row[0].text = str(i + 1)
        row[1].text = c["name"]
        row[2].text = c["usn"]
        row[3].text = str(c["total"])

    doc.save(output_path)

    return send_file(output_path, as_attachment=True)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
